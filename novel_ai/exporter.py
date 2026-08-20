"""Novel exporter: EPUB (no deps), DOCX (needs python-docx), plain text."""
from __future__ import annotations
import io
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    return text


def _paragraphs(text: str) -> List[str]:
    """Split text into paragraphs by blank lines."""
    import re
    paras = re.split(r"\n\s*\n", text.strip())
    return [p.strip() for p in paras if p.strip()]


def export_epub(title: str, author: str, chapters: List[Dict], output_path: str) -> str:
    """Export to EPUB format (pure Python, no external dependencies)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)

        # META-INF/container.xml
        container = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
            "<container version=\"1.0\" xmlns=\"urn:oasis:names:tc:opendocument:xmlns:container\">\n"
            "  <rootfiles>\n"
            "    <rootfile full-path=\"OEBPS/content.opf\" media-type=\"application/oebps-package+xml\"/>\n"
            "  </rootfiles>\n"
            "</container>\n"
        )
        zf.writestr("META-INF/container.xml", container.encode("utf-8"))

        uid = title.replace(" ", "_") + "_" + datetime.now().strftime("%Y%m%d%H%M%S")
        manifest = []
        spine = []

        manifest.append("  <item id=\"ncx\" href=\"toc.ncx\" media-type=\"application/x-dtbncx+xml\"/>\n")

        css = (
            "@namespace epub \"http://www.idpf.org/2007/ops\";\n"
            "body { font-family: serif; line-height: 1.8; margin: 5% 6%; font-size: 1em; }\n"
            "h1 { text-align: center; font-size: 1.4em; margin: 2em 0 1em 0; }\n"
            "p { text-indent: 2em; margin: 0.5em 0; line-height: 1.8; }\n"
        )
        zf.writestr("OEBPS/style.css", css.encode("utf-8"))
        manifest.append("  <item id=\"css\" href=\"style.css\" media-type=\"text/css\"/>\n")

        # Cover page
        cover = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
            "<!DOCTYPE html>\n"
            "<html xmlns=\"http://www.w3.org/1999/xhtml\">\n"
            "<head><title>Cover</title></head>\n"
            "<body>\n"
            "  <h1>" + _escape_html(title) + "</h1>\n"
            "  <p>Author: " + _escape_html(author) + "</p>\n"
            "  <p>" + datetime.now().strftime("%Y-%m-%d %H:%M") + "</p>\n"
            "</body></html>\n"
        )
        zf.writestr("OEBPS/cover.xhtml", cover.encode("utf-8"))
        manifest.append("  <item id=\"cover\" href=\"cover.xhtml\" media-type=\"application/xhtml+xml\"/>\n")
        spine.append("    <itemref idref=\"cover\"/>\n")

        # Chapters
        for i, ch in enumerate(chapters):
            ch_title = ch.get("title", "Chapter " + str(i + 1))
            ch_text = ch.get("text", "")
            paras = _paragraphs(ch_text)
            body = ""
            for p in paras:
                body += "    <p>" + _escape_html(p) + "</p>\n"

            ch_html = (
                "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
                "<!DOCTYPE html>\n"
                "<html xmlns=\"http://www.w3.org/1999/xhtml\">\n"
                "<head><title>" + _escape_html(ch_title) + "</title></head>\n"
                "<body>\n"
                "  <h1>" + _escape_html(ch_title) + "</h1>\n"
                + body +
                "</body></html>\n"
            )

            ch_file = "chapter_" + str(i + 1).zfill(3) + ".xhtml"
            zf.writestr("OEBPS/" + ch_file, ch_html.encode("utf-8"))
            manifest.append("  <item id=\"ch" + str(i + 1) + "\" href=\"" + ch_file + "\" media-type=\"application/xhtml+xml\"/>\n")
            spine.append("    <itemref idref=\"ch" + str(i + 1) + "\"/>\n")

        # content.opf
        opf = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
            "<package xmlns=\"http://www.idpf.org/2007/opf\" version=\"2.0\" unique-identifier=\"BookId\">\n"
            "  <metadata xmlns:dc=\"http://purl.org/dc/elements/1.1/\">\n"
            "    <dc:identifier id=\"BookId\">urn:uuid:" + uid + "</dc:identifier>\n"
            "    <dc:title>" + _escape_html(title) + "</dc:title>\n"
            "    <dc:creator>" + _escape_html(author) + "</dc:creator>\n"
            "    <dc:language>zh-CN</dc:language>\n"
            "    <dc:date>" + datetime.now().strftime("%Y-%m-%d") + "</dc:date>\n"
            "  </metadata>\n"
            "  <manifest>\n" + "".join(manifest) + "  </manifest>\n"
            "  <spine toc=\"ncx\">\n" + "".join(spine) + "  </spine>\n"
            "</package>\n"
        )
        zf.writestr("OEBPS/content.opf", opf.encode("utf-8"))

        # toc.ncx
        nav = []
        for i, ch in enumerate(chapters):
            ch_title = ch.get("title", "Chapter " + str(i + 1))
            ch_file = "chapter_" + str(i + 1).zfill(3) + ".xhtml"
            nav.append("    <navPoint id=\"navpoint-" + str(i + 1) + "\" playOrder=\"" + str(i + 1) + "\">\n")
            nav.append("      <navLabel><text>" + _escape_html(ch_title) + "</text></navLabel>\n")
            nav.append("      <content src=\"" + ch_file + "\"/>\n")
            nav.append("    </navPoint>\n")

        ncx = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n"
            "<ncx xmlns=\"http://www.daisy.org/z3986/2005/ncx/\" version=\"2005-1\">\n"
            "  <head>\n"
            "    <meta name=\"dtb:uid\" content=\"" + uid + "\"/>\n"
            "    <meta name=\"dtb:depth\" content=\"1\"/>\n"
            "    <meta name=\"dtb:totalPageCount\" content=\"0\"/>\n"
            "    <meta name=\"dtb:maxPageNumber\" content=\"0\"/>\n"
            "  </head>\n"
            "  <docTitle><text>" + _escape_html(title) + "</text></docTitle>\n"
            "  <navMap>\n" + "".join(nav) + "  </navMap>\n"
            "</ncx>\n"
        )
        zf.writestr("OEBPS/toc.ncx", ncx.encode("utf-8"))

    out_path = Path(output_path)
    out_path.write_bytes(buf.getvalue())
    return str(out_path.resolve())


def export_docx(title: str, author: str, chapters: List[Dict], output_path: str) -> str:
    """Export to DOCX format (needs python-docx library)."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Need python-docx: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.8

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(22)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author: " + author)
    r.font.size = Pt(14)
    doc.add_page_break()

    for i, ch in enumerate(chapters):
        ch_title = ch.get("title", "Chapter " + str(i + 1))
        ch_text = ch.get("text", "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ch_title)
        r.font.size = Pt(16)
        r.font.bold = True

        for para in _paragraphs(ch_text):
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.8
        doc.add_page_break()

    doc.save(output_path)
    return str(Path(output_path).resolve())


def export_txt(title: str, author: str, chapters: List[Dict], output_path: str) -> str:
    """Export to plain text format."""
    parts = ["=" * 20, title, "Author: " + author, "=" * 20, ""]
    for i, ch in enumerate(chapters):
        ch_title = ch.get("title", "Chapter " + str(i + 1))
        ch_text = ch.get("text", "")
        parts.append("Chapter " + str(i + 1) + ": " + ch_title)
        parts.append("")
        parts.append(ch_text)
        parts.append("")
    Path(output_path).write_text("\n".join(parts), encoding="utf-8")
    return str(Path(output_path).resolve())


def export_project(store, title: str, fmt: str, output_dir: str = None) -> Dict[str, str]:
    """Export a project to the specified format(s)."""
    chapters_raw = store.list_chapters()
    chapters = []
    for ch in chapters_raw:
        idx = ch["index"]
        p = store.chapters_dir / ("chapter_" + str(idx).zfill(3) + ".txt")
        text = p.read_text(encoding="utf-8") if p.exists() else ""
        chapters.append({"title": "Chapter " + str(idx), "text": text, "index": idx})

    out_dir = Path(output_dir) if output_dir else store.dir
    out_dir.mkdir(parents=True, exist_ok=True)
    results = {}
    formats = ["epub", "docx", "txt"] if fmt == "all" else [fmt]
    author = "AI Novel Generator"

    for f in formats:
        if f == "epub":
            results["epub"] = export_epub(title, author, chapters, str(out_dir / (title + ".epub")))
        elif f == "docx":
            results["docx"] = export_docx(title, author, chapters, str(out_dir / (title + ".docx")))
        elif f == "txt":
            results["txt"] = export_txt(title, author, chapters, str(out_dir / (title + ".txt")))

    return results